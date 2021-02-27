"""
Custom Invitations as Word Documents.
Generates a Word document with custom invitations.
"""

from pathlib import Path
import docx

# Opens text file with guest list, gets a list of them
guestFile = open(Path.cwd() / 'guests.txt')
guestList = guestFile.readlines()
# Open word document
doc = docx.Document(Path.cwd() / 'guestList.docx')
# Initialise a paragraph counter
paracount = 0

# Loop through the guestlist, creating an invite for each guest
# Note: Name and Marc are both word styles
for i in guestList:
    # Checks if it is the first in the list, or not, as the first page had
    # a new line at the start, and this keeps it consistent
    if i != guestList[0]:
        doc.add_paragraph('\nIt would be a pleasure to have \
the company of', 'Marc')
    else:
        doc.add_paragraph('It would be a pleasure to have \
the company of', 'Marc')
    # If it is not the last on the list, proceed normally, otherwise
    # it needs a new line adding itself, as this is not done automatically
    if i != guestList[-1]:
        paraObj = doc.add_paragraph(i, 'Name')
    else:
        paraObj = doc.add_paragraph(i + '\n', 'Name')
    paraObj.add_run('at 11010 Memory Lane on the Evening of', 'Marc Char')
    doc.add_paragraph('April 1st', 'Carkzis')
    # If it is not the last on the list, create a new page for the next guest
    # invitation
    lastline = doc.add_paragraph('at 7 o\'clock', 'Marc')
    if i != guestList[-1]:
        lastline.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# Save
doc.save('guestList.doc')
